import streamlit as st
import pandas as pd
from io import BytesIO, StringIO
import fitz  # PyMuPDF
from pathlib import Path
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tempfile

st.set_page_config(page_title="PVsyst Proposal Generator", layout="centered")
st.title("📊 PVsyst Proposal Generator (CSV + PDF → Filled DOCX)")

# Helper to replace unsupported characters for Word output
def safe_text(text: str) -> str:
    return text.replace("₂", "2")

# === User inputs ===
company_name = st.text_input("Enter Site Name", value="ABC Pvt Ltd")
P_nom_kWp_str = st.text_input("Capacity (power)", value='9.4')
try:
    P_nom_kWp = float(P_nom_kWp_str)
except ValueError:
    st.error("Please enter a valid number for Capacity (power).")
    st.stop()

template_file = st.file_uploader("Upload Word Template (.docx)", type=["docx"])
csv_file = st.file_uploader("Upload PVsyst Hourly CSV", type=["csv"])
pdf_file = st.file_uploader("Upload PVsyst PDF Report", type=["pdf"])

if template_file and csv_file and pdf_file:
    # --- Read CSV and find header line ---
    csv_bytes = csv_file.read()
    csv_str = csv_bytes.decode("cp1252").splitlines()

    header_line = None
    for i, line in enumerate(csv_str):
        if line.lower().startswith("date"):
            header_line = i
            break

    if header_line is None:
        st.error("Could not find header row in CSV.")
    else:
        df = pd.read_csv(
            StringIO("\n".join(csv_str)),
            skiprows=header_line,
            header=0,
            encoding="cp1252"
        )

        # Drop units row
        df = df.drop(index=0).reset_index(drop=True)
        df = df.dropna()

        # --- KPI calculations ---
        df["E_Grid"] = pd.to_numeric(df["E_Grid"], errors="coerce")
        annual_energy_mwh = df["E_Grid"].sum() / 1_000_000

        df["PR"] = pd.to_numeric(df["PR"], errors="coerce")
        pr_percent = df.loc[df["PR"] != 0, "PR"].mean() * 100

        #P_nom_kWp = 9.4  # system size
        specific_yield = (annual_energy_mwh * 1000) / P_nom_kWp

        emission_factor = 0.82
        co2_savings_tons = annual_energy_mwh * emission_factor

        # --- Show results ---
        st.subheader("KPI Summary")
        st.write(f"**Annual Energy:** {annual_energy_mwh:.2f} MWh")
        st.write(f"**Performance Ratio:** {pr_percent:.2f} %")
        st.write(f"**Specific Yield:** {specific_yield:.1f} kWh/kWp")
        st.write(f"**CO₂ Savings:** {co2_savings_tons:.1f} tons/year")

        # --- Extract Loss Diagram from PDF ---
        pdf_bytes = pdf_file.read()
        pdf_doc = fitz.open(stream=pdf_bytes, filetype="pdf")

        try:
            page_index = 4  # adjust if needed
            rect = fitz.Rect(50, 122, 550, 600)
            pix = pdf_doc[page_index].get_pixmap(clip=rect, dpi=200)

            img_bytes = BytesIO(pix.tobytes("png"))
            st.subheader("Loss Diagram")
            st.image(img_bytes, caption="Extracted Loss Diagram", use_column_width=True)

            # Save loss diagram to a temp file for docx insertion
            with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmp_img:
                tmp_img.write(img_bytes.getbuffer())
                tmp_img_path = Path(tmp_img.name)

            # --- Generate DOCX ---
            if st.button("Generate Proposal DOCX"):
                # Load template from uploaded file
                with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_tpl:
                    tmp_tpl.write(template_file.read())
                    tmp_tpl_path = Path(tmp_tpl.name)

                doc = Document(tmp_tpl_path)

                # Placeholder text mapping
                mapping = {
                    "{{ClientName}}": company_name,
                    "{{Site}}": "Navalur",
                    "{{Capacity}}": f"{P_nom_kWp} kW",
                    "{{Energy}}": f"{annual_energy_mwh:.2f} MWh",
                    "{{PR}}": f"{pr_percent:.2f} %",
                    "{{SpecificYield}}": f"{specific_yield:.1f} kWh/kWp",
                    "{{CO2Savings}}": f"{co2_savings_tons:.1f} tons/year"
                }

                # Image placeholders
                images = {
                    "{{LossDiagram}}": tmp_img_path
                }

                # Replace text in paragraphs
                for p in doc.paragraphs:
                    for key, val in mapping.items():
                        if key in p.text:
                            p.text = p.text.replace(key, safe_text(str(val)))

                # Replace text in tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for key, val in mapping.items():
                                if key in cell.text:
                                    cell.text = cell.text.replace(key, safe_text(str(val)))

                # Replace image placeholders
                for p in doc.paragraphs:
                    for key, img_path in images.items():
                        if key in p.text:
                            p.text = ""
                            run = p.add_run()
                            run.add_picture(str(img_path), width=Inches(4))
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Save to BytesIO for download
                output_docx = BytesIO()
                doc.save(output_docx)
                output_docx.seek(0)

                st.download_button(
                    label="⬇️ Download Proposal DOCX",
                    data=output_docx,
                    file_name="FilledProposal.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

        except Exception as e:
            st.error(f"Could not extract loss diagram: {e}")